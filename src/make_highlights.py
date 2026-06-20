"""Generate clean Highlights.docx for CMS submission"""
from docx import Document
from docx.shared import Pt, Inches

highlights = [
    "Uncertainty-aware perovskite stability prediction with conformal intervals.",
    "CQR reduces expected calibration error by 43-60% over standard split conformal.",
    "Three-method applicability domain separates trusted (R2=0.945) from untrusted.",
    "Leave-one-element-out over 68 elements maps extrapolation reliability per element.",
    "SR consensus validates electronegativity dominance (100% across 50 equations).",
]

# 验证字符数
print("Character count check:")
for i, h in enumerate(highlights):
    n = len(h)
    status = "OK" if n <= 85 else "OVER"
    print("  %d. [%d chars] %s - %s" % (i+1, n, status, h))

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Title
p = doc.add_paragraph()
run = p.add_run('Highlights')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Bullet points
for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(h)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.save('paper/Highlights.docx')
print("\n[SAVE] paper/Highlights.docx")
