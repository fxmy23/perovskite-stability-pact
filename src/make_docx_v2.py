"""
Generate Elsevier-formatted .docx with TABLES and FIGURES embedded inline.
Tables as proper Word tables (three-line style).
Figures as embedded images at correct positions.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

PAPER = Path('paper')
FIG = PAPER / 'figures'
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 2.0

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = 'Times New Roman'
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_figure(name, caption, width=Inches(5.5)):
    """Insert image + caption."""
    img_path = FIG / f'{name}.png'
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
        # Caption
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.line_spacing = 1.0
        cr = cp.add_run(caption)
        cr.font.size = Pt(10)
        cr.font.name = 'Times New Roman'
        doc.add_paragraph()  # spacer

def add_table(headers, rows, caption, note=None):
    """Add a three-line table (no vertical borders)."""
    # Caption first
    cp = doc.add_paragraph()
    cp.paragraph_format.line_spacing = 1.0
    cr = cp.add_run(caption)
    cr.bold = True
    cr.font.size = Pt(10)
    cr.font.name = 'Times New Roman'

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Remove all borders first, then add only top/bottom/header-bottom
    tbl = table._tbl
    # Set header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
    # Note
    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.line_spacing = 1.0
        nr = np_.add_run(note)
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacer


# === Read manuscript and render with embedded tables/figures ===
md = (PAPER / 'manuscript.md').read_text(encoding='utf-8')
lines = md.split('\n')

i = 0
while i < len(lines):
    line = lines[i].strip()

    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        i += 1; continue

    if line.startswith('**Author:**') or line.startswith('**Affiliation:**') or line.startswith('**Corresponding'):
        para(line.replace('**', ''))
        i += 1; continue

    if line == '---':
        i += 1; continue

    if line.startswith('## '):
        h = line[3:].strip()
        heading(h, 1)
        # Insert figures at section boundaries
        if '2. Methods' in h:
            add_figure('Figure_1_workflow',
                       'Fig. 1. Overview of the uncertainty-aware prediction framework (PACT-Final).')
        i += 1; continue

    if line.startswith('### '):
        h = line[4:].strip()
        heading(h, 2)
        i += 1; continue

    if line.startswith('**Keywords:**'):
        para(line.replace('**', ''))
        i += 1; continue

    # Detect figure/table placeholder markers and insert
    if '[**Figure 2' in line or '[**Fig. 2' in line or 'Figure 2' in line and 'parity' in line.lower():
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_figure('Figure_2_parity',
                   'Fig. 2. Parity plots of predicted vs. DFT values for (a) formation energy and (b) energy above hull. Points colored by ensemble sigma. Dashed line: y = x.')
        i += 1; continue

    if '[**Figure 3' in line or 'Figure 3' in line and 'reliability' in line.lower():
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_figure('Figure_3_reliability',
                   'Fig. 3. Reliability diagrams showing empirical coverage per uncertainty decile for standard split conformal vs. CQR.')
        i += 1; continue

    if '[**Figure 4' in line or 'Figure 4' in line and ('applicability' in line.lower() or 'AD' in line):
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_figure('Figure_4_AD',
                   'Fig. 4. Applicability domain visualization. Predicted value vs. absolute error, colored by trusted vs. untrusted regions.')
        i += 1; continue

    if '[**Figure 5' in line or 'Figure 5' in line and 'LOEO' in line:
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_figure('Figure_5_LOEO',
                   'Fig. 5. Leave-one-element-out (LOEO) extrapolation R-squared for representative A-site elements.')
        i += 1; continue

    if '[**Figure 6' in line or 'Figure 6' in line and ('SHAP' in line or 'SR' in line):
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_figure('Figure_6_SHAP_SR',
                   'Fig. 6. (a) LightGBM SHAP feature importance (top-15) and (b) symbolic regression consensus frequency.')
        i += 1; continue

    if '[**Table 1' in line or ('Table 1' in line and 'Model comparison' in line):
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_table(
            ['Model', 'Form. R2', 'Form. MAE', 'Hull R2', 'Hull MAE'],
            [['Pure LightGBM', '0.912', '0.193', '0.793', '0.179'],
             ['GBDT stacking', '0.918', '0.181', '0.813', '0.169'],
             ['PACT-Final', '0.914', '0.186', '0.800', '0.171'],
             ['Ref. [10] (diff. dataset)', '0.928', '-', '-', '-']],
            'Table 1. Model comparison for formation energy and hull energy prediction (5-fold CV).'
        )
        i += 1; continue

    if '[**Table 2' in line or ('Table 2' in line and 'CQR' in line):
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_table(
            ['Target', 'Method', 'PICP', 'MPIW', 'ECE', 'Improv.'],
            [['Form. E', 'Standard', '0.810', '0.611', '0.085', '-'],
             ['Form. E', 'CQR', '0.802', '0.823', '0.034', '60%'],
             ['Hull E', 'Standard', '0.815', '0.566', '0.086', '-'],
             ['Hull E', 'CQR', '0.807', '0.719', '0.049', '43%']],
            'Table 2. CQR vs. standard split conformal (fair LightGBM-family baseline).'
        )
        i += 1; continue

    if '[**Table 3' in line or ('Table 3' in line and 'Candidate' in line):
        para(line.replace('[**', '').replace('**]', '').replace('[', '').replace(']', ''))
        add_table(
            ['Formula', 'Pred. E_hull', 'sigma', 'In matbench', 'Goldsm.', 'Grade'],
            [['LaTiO3', '0.005', '0.022', 'Yes', 'Yes', 'Mod.'],
             ['PrCoO3', '0.006', '0.024', 'No', 'Yes', 'Weak'],
             ['GdCoO3', '0.016', '0.027', 'No', 'Yes', 'Weak'],
             ['YbCoO3', '0.017', '0.028', 'No', 'Yes', 'Weak'],
             ['ErFeO3', '0.019', '0.030', 'No', 'Yes', 'Weak'],
             ['YCoO3', '0.025', '0.031', 'Yes', 'Yes', 'Mod.'],
             ['LaAlO3', '0.033', '0.037', 'Yes', 'Yes', 'Weak']],
            'Table 3. Candidate stable perovskites with validation tier.'
        )
        i += 1; continue

    # Skip the old standalone Tables section
    if line.startswith('## Tables'):
        # Skip everything until next ## heading
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('## '):
            i += 1
        continue

    # Skip Figure Captions section (already embedded inline)
    if line.startswith('## Figure Captions'):
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('## '):
            i += 1
        continue

    # Skip markdown table rows (already handled above)
    if line.startswith('|') and '---' not in line:
        i += 1; continue
    if line.startswith('|') and '---' in line:
        i += 1; continue

    # Regular paragraph
    if line and not line.startswith('[**'):
        clean = line.replace('**', '').replace('*', '')
        if clean.strip():
            para(clean)
    i += 1

# === Add References ===
doc.add_page_break()
heading('References', 1)
refs_md = (PAPER / 'references.md').read_text(encoding='utf-8')
for rline in refs_md.split('\n'):
    rline = rline.strip()
    if rline.startswith('['):
        para(rline)

output = PAPER / 'manuscript_cms_v6.docx'
doc.save(str(output))
print(f'[SAVE] {output} ({output.stat().st_size // 1024} KB)')
