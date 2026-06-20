"""
Update citation numbers + reference list in manuscript.
- Maps old [X] to new [X] per dedup table
- Deletes [5] (Rajan) and merges [20,35]→[18]
- Rebuilds reference list with 30 clean entries
- Operates paragraph-by-paragraph, before References section only
- Preserves run formatting (uses first run's font for rebuilt paragraph)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import re
import copy

INPUT = 'paper/manuscript_cms_final_v3.docx'
OUTPUT = 'paper/manuscript_cms_final_v4.docx'

# Old -> new mapping (None = deleted)
MAPPING = {
    1: 1, 2: 2, 3: 3, 4: 4,
    5: None,           # Rajan deleted
    6: 5, 7: 6, 8: 7, 9: 8,
    10: 9, 11: 10,
    12: None,          # dup of 10
    13: 11, 14: 12, 15: 13, 16: 14,
    17: 15, 18: 16, 19: 17, 20: 18,
    21: 19, 22: 20, 23: 21, 24: 22,
    25: 23, 26: 24, 27: 25, 28: 26,
    29: 27,
    30: None, 31: None,  # dup of 21/29
    32: 28,
    33: None, 34: None,  # dup of 4/2
    35: None, 36: None,  # dup of 20/22
    37: 29, 38: 30,
    39: None, 40: None,  # dup of 1/15
}


def remap_citation_group(nums):
    """Remap a list of old numbers to sorted unique new numbers (drop None)."""
    new_nums = []
    for n in nums:
        new = MAPPING.get(n)
        if new is not None and new not in new_nums:
            new_nums.append(new)
    new_nums.sort()
    return new_nums


def rewrite_citations_in_text(text):
    """Replace every [..] citation in text using MAPPING. Handles ranges and lists."""
    def repl(m):
        raw = m.group(1)
        # Parse numbers (handle comma and en-dash/ascii dash ranges)
        nums = []
        for part in raw.split(','):
            part = part.strip()
            if '\u2013' in part or '-' in part:
                # range
                seg = re.split(r'[\u2013\-]', part)
                a, b = int(seg[0]), int(seg[1])
                nums.extend(range(a, b + 1))
            else:
                nums.append(int(part))
        new_nums = remap_citation_group(nums)
        if not new_nums:
            return ''  # citation deleted entirely
        # Render as comma list (avoid ranges to keep it simple and explicit)
        return '[' + ','.join(str(n) for n in new_nums) + ']'

    # Match citations like [5], [3,6], [1-4], [1–4], [20,35], [29–31]
    return re.sub(r'\[(\d+(?:\s*[,\u2013\-]\s*\d+)*)\]', repl, text)


# === Clean up residual double spaces and trailing space before punctuation ===
def cleanup(text):
    # 'word [5], next' with deleted [5] => 'word , next' -> fix
    text = re.sub(r'\s+,', ',', text)
    text = re.sub(r'\s+\.', '.', text)
    text = re.sub(r'\s{2,}', ' ', text)
    # 'pymatgen  [20,35]' had two spaces originally
    return text


def rebuild_paragraph(p, new_text):
    """Replace paragraph text while preserving the first run's formatting."""
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    # capture format
    font_name = first.font.name
    font_size = first.font.size
    bold = first.font.bold
    italic = first.font.italic
    color = first.font.color.rgb if first.font.color and first.font.color.type else None
    # clear all runs
    for run in list(p.runs):
        run.text = ''
    # remove extra runs from xml
    rPr = first._element.find(qn('w:rPr'))
    rPr_copy = copy.deepcopy(rPr) if rPr is not None else None
    # delete all run elements
    for r in list(p._element.findall(qn('w:r'))):
        p._element.remove(r)
    # add single new run
    new_run = p.add_run(new_text)
    if font_name:
        new_run.font.name = font_name
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        new_run.font.size = font_size
    if bold is not None:
        new_run.font.bold = bold
    if italic is not None:
        new_run.font.italic = italic
    if color is not None:
        new_run.font.color.rgb = color


# === New reference list (30 entries) ===
NEW_REFS = """[1] R.K. Vasudevan, K. Choudhary, A. Mehta, R. Pachter, S.V. Kalinin, M. Chi, Physics-informed machine learning for materials discovery, MRS Commun. 11 (2021) 611\u2013622. https://doi.org/10.1557/s43579-021-00050-5

[2] S. Yip (Ed.), Handbook of Materials Modeling, Springer, 2020. https://doi.org/10.1007/978-3-319-44680-6

[3] Q. Tao, L. Xu, P. Xu, Y. Liu, T. Wen, S. Du, et al., Machine learning for perovskite materials design and discovery, npj Comput. Mater. 7 (2021) 23. https://doi.org/10.1038/s41524-021-00495-8

[4] K.T. Butler, D.W. Davies, H. Cartwright, O. Isayev, A. Walsh, Machine learning for molecular and materials science, Nature 559 (2018) 547\u2013555. https://doi.org/10.1038/s41586-018-0337-z

[5] A.A. Emery, C. Wolverton, High-throughput DFT calculations of formation energy, stability and oxygen vacancy formation energy of ABO3 perovskites, Sci. Data 4 (2017) 170153. https://doi.org/10.1038/sdata.2017.153

[6] Z. Deng, K. Fang, C. Guo, Z. Gong, H. Yue, H. Zhang, F.E.H. Tay, Prediction of ABX3 perovskite formation energy using machine learning, Materials 18 (2025) 2927. https://doi.org/10.3390/ma18132927

[7] R. Ouyang, S. Curtarolo, E. Ahmetcik, M. Scheffler, L.M. Ghiringhelli, SISSO: A compressed-sensing method for identifying the best physics-inspired descriptor, Phys. Rev. Mater. 2 (2018) 083802. https://doi.org/10.1103/PhysRevMaterials.2.083802

[8] T. Xie, J.C. Grossman, Crystal graph convolutional neural networks for an accurate and interpretable prediction of material properties, Phys. Rev. Lett. 120 (2018) 145301. https://doi.org/10.1103/PhysRevLett.120.145301

[9] T.I. Netzeva, A.P. Worth, T. Aldenberg, R. Benigni, M.T.D. Cronin, P. Gramatica, J.S. Jaworska, S. Kahn, G. Klopman, C.A. Marchant, et al., Current status of methods for defining the applicability domain of (quantitative) structure\u2013activity relationships, ATLA 33 (2005) 155\u2013173. https://doi.org/10.1177/026119290503300208

[10] F. Sahigara, D. Mansouri, D. Ballabio, A. Mauri, V. Consonni, R. Todeschini, Comparison of different approaches to define the applicability domain of QSAR models, Molecules 17 (2012) 4791\u20134810. https://doi.org/10.3390/molecules17054791

[11] V. Vovk, A. Gammerman, G. Shafer, Algorithmic Learning in a Random World, Springer, New York, 2005.

[12] Y. Romano, E. Patterson, E.J. Cand\u00e8s, Conformalized quantile regression, in: Advances in Neural Information Processing Systems (NeurIPS), vol. 32, 2019, pp. 3543\u20133553. https://doi.org/10.48550/arXiv.1905.03222

[13] A. Ziletti, D. Kumar, M. Scheffler, L.M. Ghiringhelli, Insightful classification of crystal structures by deep learning, Nat. Commun. 9 (2018) 2775. https://doi.org/10.1038/s41467-018-05169-6

[14] M. Cranmer, Interpretable machine learning for science with PySR and SymbolicRegression.jl, arXiv:2305.01582, 2023. https://doi.org/10.48550/arXiv.2305.01582

[15] L. Ward, M. Dunston, D. Donadio, C. Wolverton, Matminer: An open source toolkit for materials data mining, Comput. Mater. Sci. 152 (2018) 60\u201369. https://doi.org/10.1016/j.commatsci.2018.05.018

[16] J.E. Saal, S. Kirklin, M. Aykol, B. Meredig, C. Wolverton, Materials design and discovery with high-throughput density functional theory: The Open Quantum Materials Database (OQMD), JOM 65 (2013) 1501\u20131509. https://doi.org/10.1007/s11837-013-0755-4

[17] L. Ward, A. Agrawal, A. Choudhary, C. Wolverton, A general-purpose machine learning framework for predicting properties of inorganic materials, npj Comput. Mater. 2 (2016) 16028. https://doi.org/10.1038/npjcompumats.2016.28

[18] S.P. Ong, W.D. Richards, A. Jain, G. Hautier, M. Kocher, S. Cholia, D. Gunter, V.L. Chevrier, K.A. Persson, G. Ceder, Python Materials Genomics (pymatgen): A robust, open-source python library for materials analysis, Comput. Mater. Sci. 68 (2013) 314\u2013319. https://doi.org/10.1016/j.commatsci.2012.10.028

[19] V.M. Goldschmidt, Die Gesetze der Krystallochemie, Naturwissenschaften 14 (1926) 477\u2013485. https://doi.org/10.1007/BF01507527

[20] F. Pedregosa, G. Varoquaux, A. Gramfort, V. Michel, B. Thirion, O. Grisel, M. Blondel, P. Prettenhofer, R. Weiss, V. Dubourg, J. Vanderplas, A. Passos, D. Cournapeau, M. Brucher, M. Perrot, \u00c9. Duchesnay, Scikit-learn: Machine learning in Python, J. Mach. Learn. Res. 12 (2011) 2825\u20132830. https://doi.org/10.5555/1953048.2078195

[21] T. Akiba, S. Sano, T. Yanase, T. Ohta, M. Koyama, Optuna: A next-generation hyperparameter optimization framework, in: Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining, 2019, pp. 2623\u20132631. https://doi.org/10.1145/3292500.3330701

[22] T. Hastie, R. Tibshirani, J. Friedman, The Elements of Statistical Learning, 2nd ed., Springer, 2009. https://doi.org/10.1007/978-0-387-84858-7

[23] B. Efron, R.J. Tibshirani, An Introduction to the Bootstrap, Chapman & Hall/CRC, 1993. https://doi.org/10.1201/9780429246593

[24] F. Wilcoxon, Individual comparisons by ranking methods, Biom. Bull. 1 (1945) 80\u201383. https://doi.org/10.2307/3001968

[25] S.M. Lundberg, S.-I. Lee, A unified approach to interpreting model predictions, in: Advances in Neural Information Processing Systems (NeurIPS), vol. 30, 2017. https://doi.org/10.48550/arXiv.1705.07874

[26] C.J. Bartel, C. Sutton, B.R. Goldsmith, C. Ophus, R. Ouyang, C.B. Musgrave, L.M. Ghiringhelli, M. Scheffler, Physical descriptor for the Gibbs energy of inorganic crystalline solids and temperature-dependent materials chemistry, Nat. Commun. 11 (2020) 6100. https://doi.org/10.1038/s41467-020-19908-0

[27] C.J. Bartel, C. Sutton, B.R. Goldsmith, R. Ouyang, C.B. Musgrave, L.M. Ghiringhelli, M. Scheffler, New tolerance factor to predict the stability of perovskite oxides and halides, Sci. Adv. 5 (2019) eaav0693. https://doi.org/10.1126/sciadv.aav0693

[28] A. Jain, S.P. Ong, G. Hautier, W. Chen, W.D. Richards, S. Dacek, S. Cholia, D. Gunter, D. Skinner, G. Ceder, K.A. Persson, Commentary: The Materials Project: A materials genome approach to accelerating materials innovation, APL Mater. 1 (2013) 011002. https://doi.org/10.1063/1.4812323

[29] J.D. Hunter, Matplotlib: A 2D graphics environment, Comput. Sci. Eng. 9 (2007) 90\u201395. https://doi.org/10.1109/MCSE.2007.55

[30] S. Curtarolo, W. Setyawan, G.L.W. Hart, M. Jahnatek, R.V. Chepulskii, R.H. Taylor, S. Wang, J. Xue, K. Yang, O. Levy, M.J. Mehl, H.T. Stokes, D.O. Demchenko, D. Morgan, AFLOW: An automatic framework for high-throughput materials discovery, Comput. Mater. Sci. 58 (2012) 218\u2013226. https://doi.org/10.1016/j.commatsci.2012.02.005"""


def main():
    doc = Document(INPUT)

    # Find References heading
    ref_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'References' and (p.style.name or '').startswith('Heading'):
            ref_start = i
            break
    print(f'References starts at paragraph {ref_start}')

    # === Phase 1: rewrite citations in body (before References) ===
    changes = 0
    for i in range(ref_start):
        p = doc.paragraphs[i]
        old = p.text
        if '[' not in old:
            continue
        new = rewrite_citations_in_text(old)
        new = cleanup(new)
        if new != old:
            rebuild_paragraph(p, new)
            changes += 1
    print(f'Phase 1: rewrote citations in {changes} body paragraphs')

    # === Phase 2: replace reference list ===
    # Delete all paragraphs from References heading onward (we'll re-add)
    body = doc.element.body
    # collect paragraph elements at/after ref_start
    ref_para_elems = []
    para_count = 0
    for child in list(body.iterchildren()):
        if child.tag == qn('w:p'):
            if para_count >= ref_start:
                ref_para_elems.append(child)
            para_count += 1

    # remove them
    for elem in ref_para_elems:
        body.remove(elem)
    print(f'Phase 2: removed {len(ref_para_elems)} old reference paragraphs')

    # Add new References heading + entries
    h = doc.add_heading('References', level=1)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

    for line in NEW_REFS.split('\n\n'):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)  # hanging indent
        p.paragraph_format.left_indent = Cm(0.75)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    print(f'Phase 2: added 30 new reference entries')

    doc.save(OUTPUT)
    print(f'[OK] saved: {OUTPUT}')


from docx.shared import Cm
main()
