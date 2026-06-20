"""
生成完整 §3 (4 节版), 保留 v3 原版所有内容 + 修正错误数字 + 合并 6→4 节.
所有数字 100% 核实于 results/metrics/ 和 results/discovery/.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Times New Roman'
    p.add_run(text)
    return p


def add_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_table(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, row in enumerate(table.rows):
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    return table


# ========== §3 标题 ==========
add_heading('3. Results and Discussion', level=1)

# ============================================================
# §3.1 Point prediction performance
# ============================================================
add_heading('3.1 Point prediction performance', level=1)

# Table 1 caption
add_caption('Table 1. Model comparison for formation energy and hull energy prediction (5-fold CV). Ref. [7] uses a distinct Materials-Project-derived ABX3 dataset and is included for context only.')

# Table 1
add_table(
    headers=['Model', 'Form. R²', 'Form. MAE', 'Hull R²', 'Hull MAE'],
    rows=[
        ['Pure LightGBM', '0.915', '0.193', '0.807', '0.179'],
        ['GBDT stacking', '0.918', '0.181', '0.813', '0.169'],
        ['PACT-Final', '0.914', '0.186', '0.800', '0.171'],
        ['Ref. [7] (diff. dataset)', '0.928', '—', '—', '—'],
    ],
    col_widths=[1.8, 1.0, 1.0, 1.0, 1.0]
)

# Fig. 4 caption
add_caption('Fig. 4. Parity plots of predicted vs. DFT values for (a) formation energy and (b) energy above hull. Points colored by ensemble σ. Dashed line: y = x.')

# 主段落 (保留原版 + 修正 0.358 + 加入解释)
add_para(
    "The framework achieves R² = 0.914 [0.907, 0.922] for formation energy and "
    "R² = 0.800 [0.781, 0.816] for hull energy (5-fold CV, bootstrap 95% CI), with "
    "MAE of 0.186 and 0.171 eV/atom respectively (Table 1). For context, this "
    "formation-energy R² is comparable to the gradient-boosting benchmark of Emery "
    "and Wolverton (R² ≈ 0.91 on the same wolverton_oxides dataset [6]) and to the "
    "0.928 reported by Deng et al. [7] on a distinct Materials-Project-derived ABX3 "
    "dataset using SHAP [27]-guided feature selection. Because the two benchmarks "
    "rely on different DFT reference frames (our wolverton_oxides dataset and "
    "Matbench Perovskites differ structurally and energetically), direct cross-"
    "dataset numerical comparison is not meaningful; nonetheless, the point-estimate "
    "accuracy places our model in the same tier as recent descriptor-ML benchmarks. "
    "Crucially, our differentiation is not in R² but in the uncertainty and "
    "applicability components that follow. Hull-energy prediction is inherently "
    "harder (R² = 0.800): E_hull depends on all competing phases, a global property "
    "not fully captured by single-compound descriptors. The stacking ensemble "
    "(LightGBM + XGBoost + HistGradientBoosting) marginally improves over single "
    "LightGBM (+0.003 R²), consistent with the high correlation (0.99) among GBDT "
    "family members limiting ensemble diversity gains. We note that PACT-Final "
    "(R² = 0.914) is marginally below the GBDT stacking alone (R² = 0.918) for "
    "formation energy, because the physics baseline serves as an interpretability "
    "anchor rather than an accuracy booster (§2.3); the 0.004 R² difference falls "
    "within the bootstrap CI. The advantage of PACT-Final over pure stacking lies "
    "not in point accuracy but in the calibrated CQR interval and applicability "
    "domain reported in §3.2–3.3."
)

add_para(
    "The physics baseline alone (KernelRidge on 14 physics features) explains "
    "R² = 0.773 of formation-energy variance—confirming that physically meaningful "
    "descriptors (tolerance factor, electronegativity, ionic radii) capture the "
    "majority of the signal, consistent with the ionic-model understanding [28] of "
    "perovskite stability. The stacking residual adds the non-linear/secondary-"
    "descriptor structure.",
    bold_prefix="Material interpretation. "
)

add_para(
    "The 8.0% of samples with absolute error exceeding 0.5 eV/atom (394/4914) "
    "cluster at the extreme of the formation-energy distribution: their mean "
    "E_f = −1.384 eV/atom is less negative than the global mean (−1.663), "
    "indicating the model struggles most with compounds of intermediate stability—"
    "neither strongly stable nor clearly unstable—where small descriptor changes "
    "produce large energy shifts. Conversely, the most accurately predicted "
    "compounds are strongly stable perovskites (E_f < −2.5 eV/atom, e.g., "
    "alkaline-earth-based oxides) whose formation energy is dominated by well-"
    "captured ionic contributions. The ensemble uncertainty σ correlates positively "
    "with absolute error in both targets (formation: r = 0.345, p ≈ 1.3 × 10⁻¹³⁷; "
    "hull: r = 0.355, p ≈ 1.4 × 10⁻¹⁴⁵), confirming that the model's internal "
    "disagreement ranks prediction difficulty correctly. Notably, 4 of the 5 worst "
    "predictions occur at σ < 0.10 eV/atom, suggesting the ensemble underestimates "
    "uncertainty for certain extrapolative samples; this residual miscalibration "
    "motivates the conditional (CQR) intervals in §3.2, which do not rely on σ alone.",
    bold_prefix="Error analysis. "
)

# ============================================================
# §3.2 Uncertainty quantification: CQR conditional coverage
# ============================================================
add_heading('3.2 Uncertainty quantification: CQR conditional coverage', level=1)

# Table 2 caption
add_caption('Table 2. CQR vs. standard split conformal (fair LightGBM-family baseline). PICP: prediction interval coverage probability; MPIW: mean prediction interval width; ECE: expected calibration error.')

# Table 2
add_table(
    headers=['Target', 'Method', 'PICP', 'MPIW', 'ECE', 'Improv.'],
    rows=[
        ['Form. E', 'Standard', '0.810', '0.611', '0.085', '—'],
        ['Form. E', 'CQR', '0.802', '0.823', '0.034', '60%'],
        ['Hull E', 'Standard', '0.815', '0.566', '0.086', '—'],
        ['Hull E', 'CQR', '0.807', '0.719', '0.049', '43%'],
    ],
    col_widths=[1.0, 1.0, 0.9, 0.9, 0.9, 0.9]
)

# Fig. 5 caption
add_caption('Fig. 5. Reliability diagrams showing empirical coverage per uncertainty decile for standard split conformal vs. CQR.')

# 主段落 (保留原版全部内容)
add_para(
    "CQR delivers intervals satisfying the marginal coverage guarantee (PICP = "
    "0.802 / 0.807, both ≥ 0.80 nominal) while reducing ECE by 43–60% relative to "
    "standard split conformal under a fair (same LightGBM-family) baseline (Table 2): "
    "formation-energy ECE drops from 0.085 to 0.034 (−60%), hull-energy from 0.086 "
    "to 0.049 (−43%). This means conditional coverage is markedly more uniform—high-"
    "uncertainty (extrapolation) samples no longer suffer severe under-coverage, and "
    "low-uncertainty (interpolation) samples receive appropriately tight intervals."
)

add_para(
    "The heteroscedasticity is empirically visible: the mean interval width for the "
    "highest-σ tercile (1.03 eV/atom for formation energy) is ~1.5× that of the "
    "lowest-σ tercile (0.68), whereas standard conformal assigns uniform width "
    "(0.66 to both). For high-throughput [38] screening, this sample-adaptive "
    "behavior directly translates to fewer false-positive stable predictions among "
    "uncertain candidates. This behavior is further visualized in Fig. 2, where the "
    "CQR interval width grows in tandem with the absolute error across samples sorted "
    "by ensemble σ, whereas the standard interval remains uniform."
)

add_para(
    "The σ–error Pearson correlation (r = 0.345 for formation energy, p = 1.3 × "
    "10⁻¹³⁷) confirms that ensemble disagreement tracks actual error: when the GBDT "
    "models disagree, the prediction is genuinely less reliable. This physical "
    "meaningfulness of the uncertainty estimate is a prerequisite for AD-based "
    "screening (§3.3).",
    bold_prefix="Material interpretation. "
)

# ============================================================
# §3.3 Applicability domain and extrapolation
# ============================================================
add_heading('3.3 Applicability domain and extrapolation', level=1)

# Fig. 6 caption (原 Fig. 4)
add_caption('Fig. 6. Applicability domain visualization. Predicted value vs. absolute error, colored by trusted vs. untrusted regions.')

# AD 主段落 (保留原版)
add_para(
    "The ensemble-σ AD criterion partitions samples into a trusted region (σ < "
    "median) achieving R² = 0.945 (formation energy) versus R² = 0.886 in the "
    "untrusted region—a 0.059 R² gap that validates the criterion. The k-NN distance "
    "and PCA leverage criteria yield trusted-region R² of 0.915 and 0.916 "
    "respectively, with 50–53% agreement with the σ criterion. The three methods "
    "provide complementary, convergent evidence for the model's reliability boundary "
    "rather than a single arbitrary threshold (Fig. 3)."
)

add_para(
    "For experimentalists using the model to prioritize DFT verification, restricting "
    "to the trusted region roughly halves the error variance, directly reducing "
    "wasted computation on unreliable predictions.",
    bold_prefix="Material interpretation. "
)

# LOEO 段落 (修正 68→73, 0.70→0.74, 删 B/Ac 负 R²)
add_para(
    "To probe extrapolation behavior more directly, we perform a leave-one-element-"
    "out (LOEO) evaluation: for each of the 73 elements appearing at the A or B "
    "site, we retrain on all samples not containing that element and evaluate on the "
    "held-out element's compounds. The mean per-element R² across the 73 element-"
    "wise test sets is 0.739 (median 0.789) for formation energy, substantially "
    "lower than the in-domain R² of 0.914—as expected for compositional "
    "extrapolation—but well above zero for all elements, with the lowest-performing "
    "elements being Pb (0.224), Al (0.246), and Os (0.275). No element yields a "
    "negative R². Fig. 7 displays the full element-wise distribution."
)

# Fig. 7 caption (原 Fig. 5)
add_caption('Fig. 7. Leave-one-element-out (LOEO) extrapolation R-squared for all 73 A-site elements.')

# LOEO material interpretation (修正 R² 数字)
add_para(
    "The extrapolation pattern is chemically coherent. The best-extrapolated A-site "
    "elements are large, highly electropositive cations—alkali metals (Rb R²=0.96, "
    "K 0.93, Na 0.94, Cs 0.92), alkaline earths (Ba 0.94), and lanthanides (Pr 0.95, "
    "La 0.91, Gd 0.92, Ho 0.93, Dy 0.94). These elements share consistent +2 or +3 "
    "oxidation states, large ionic radii, and well-defined coordination chemistry, so "
    "the descriptor space around them is densely sampled and their A-site "
    "contributions to formation energy are nearly linear in ionic radius and "
    "electronegativity. In contrast, the worst-extrapolated elements occupy sparsely "
    "populated or chemically anomalous regions of the descriptor space: Pb (0.22), "
    "Al (0.25), and Os (0.27) have very few isochemical neighbors in the training "
    "set, and their bonding character (covalent/metallic rather than ionic for Al; "
    "heavy 5d chemistry for Os) deviates from the perovskite norm. This element-"
    "specific failure map directly informs screening: predictions involving "
    "underrepresented heavy p-block or atypical main-group A-site cations should be "
    "treated as unreliable regardless of the point estimate, whereas screening among "
    "alkali/alkaline-earth/lanthanide chemistries is expected to transfer well.",
    bold_prefix="Material interpretation. "
)

# ============================================================
# §3.4 Interpretability and candidate screening
# ============================================================
add_heading('3.4 Interpretability and candidate screening', level=1)

# SR 共识段落 (保留原版, SNR 已在 §2.7 定义)
add_para(
    "Across 50 independent SR equations (10 seeds × 5 folds), A-site "
    "electronegativity appears in 100% of equations, with B-site group number (64%) "
    "and B-site electronegativity (60%) the next most conserved. This consensus "
    "validates the known dominance of electronegativity in governing perovskite "
    "formation energy (consistent with ionic/covalent bonding theory) and provides "
    "an interpretable cross-check that the ML model has captured physically "
    "meaningful structure."
)

# Physical direction 段落 (保留原版全部)
add_para(
    "To verify that both the ML model and the symbolic regression have captured "
    "chemically correct structure–property relationships, we examine the sign and "
    "magnitude of each top descriptor's correlation with formation energy (Table S3). "
    "The strongest single correlate is A-site electronegativity (Pearson r = +0.537): "
    "more electronegative A-site cations (e.g., transition metals at the A site) "
    "yield less negative (less stable) formation energies, consistent with the fact "
    "that highly electropositive cations (alkali, alkaline-earth, lanthanide) form "
    "stronger ionic bonds with oxygen and stabilize the perovskite lattice. "
    "Conversely, Magpie X_std (the compositional standard deviation of "
    "electronegativity) is the strongest negative correlate (r = −0.538): greater "
    "electronegativity contrast among A/B/O sites produces more negative formation "
    "energies, reflecting the thermodynamic driving force of charge transfer in "
    "ionic bonding. The tolerance factor t shows a moderate negative correlation "
    "(r = −0.278): as t increases toward and beyond unity (A-site too large for the "
    "cage), the structure becomes geometrically strained and formation energy rises—"
    "consistent with Goldschmidt's criterion. B-site group number (r = +0.297) and "
    "B-site electronegativity (r = +0.291) are positively correlated: higher B-site "
    "electronegativity (late transition metals) reduces the ionic character of the "
    "B–O bond, destabilizing the perovskite relative to competing phases. Crucially, "
    "the SHAP-derived ranking and the SR consensus agree on the identity of the top "
    "descriptors even though they operate in different spaces (statistical vs. "
    "symbolic), and the Pearson signs are consistent with established perovskite "
    "chemistry [29] [30,31], confirming the model is not exploiting spurious "
    "correlations.",
    bold_prefix="Physical direction of key features. "
)

# Fig. 8 caption (原 Fig. 6)
add_caption('Fig. 8. (a) LightGBM SHAP feature importance (top-15) and (b) symbolic regression consensus frequency.')

# SR applicability 段落 (修正 SNR 2.61→3.40, 0.94→1.39)
add_para(
    "The SR applicability is target-dependent: SR succeeds on formation energy "
    "(standalone R² ≈ 0.44 across 10 seeds, SNR = 3.40) but fails on hull energy "
    "(SNR = 1.39, with the SR solution collapsing to a constant in several folds). "
    "We trace this to the differing physical nature of the two targets—formation "
    "energy is dominated by single-compound bond chemistry (near-linear, well-"
    "expressed by basic operators), whereas hull energy depends on competition with "
    "all other phases (strongly non-linear, non-monotonic in geometric descriptors). "
    "This yields an empirical SR applicability criterion (SNR ≥ 2), which we caveat "
    "as based on only two targets and not yet universal."
)

# 候选段落 + Table 3
add_caption('Table 3. Candidate stable perovskites with validation tier.')

add_table(
    headers=['Formula', 'Pred. E_hull', 'σ', 'In matbench', 'Goldsm.', 'Grade'],
    rows=[
        ['LaTiO3', '0.005', '0.022', 'Yes', 'Yes', 'Mod.'],
        ['PrCoO3', '0.006', '0.024', 'No', 'Yes', 'Weak'],
        ['GdCoO3', '0.016', '0.027', 'No', 'Yes', 'Weak'],
        ['YbCoO3', '0.017', '0.028', 'No', 'Yes', 'Weak'],
        ['DyCoO3', '0.019', '0.029', 'No', 'Yes', 'Weak'],
        ['ErFeO3', '0.019', '0.030', 'No', 'Yes', 'Weak'],
        ['YCoO3', '0.025', '0.031', 'Yes', 'Yes', 'Mod.'],
        ['SrNpO3', '0.025', '0.034', 'No', 'No', 'Weak'],
        ['LaAlO3', '0.033', '0.037', 'Yes', 'Yes', 'Weak'],
    ],
    col_widths=[1.2, 1.1, 0.8, 1.0, 0.9, 0.8]
)

add_para(
    "Applying the framework to virtual A–B combinations yields 9 candidate stable "
    "perovskites (predicted E_hull < 0.05 eV/atom within the trusted AD region). "
    "Three-tier validation provides indirect evidence: 3 candidates (LaTiO3, YCoO3, "
    "LaAlO3) appear in the independent matbench_perovskites DFT database (confirming "
    "they have been studied); 8 of 9 fall within the Goldschmidt synthesizability "
    "zone (t ∈ [0.8, 1.05], μ ∈ [0.4, 0.9]); and leave-element-out robustness testing "
    "confirms the predictions are not artifacts of single-element memorization for "
    "the moderate-grade candidates. We emphasize that database presence confirms "
    "study, not numerical accuracy of our predictions (due to the reference-frame "
    "mismatch noted in §2.1), and that DFT verification remains necessary."
)

doc.save('paper/section_3_revised_v2.docx')
print('[OK] saved: paper/section_3_revised_v2.docx')
print()
print('=== 修正的数字清单 ===')
print('  §3.1 Table 1 Pure LightGBM Hull R²: 0.793 → 0.807 (实测)')
print('  §3.1 删除 Pearson r=0.358, 改为定性说明')
print('  §3.1 加入 PACT-Final vs stacking 的解释段')
print('  §3.1 加入 σ-error 相关性 r=0.345/0.355')
print('  §3.3 LOEO: 68→73 元素, 0.70→0.739 R², 删 B/Ac 负 R²')
print('  §3.3 LOEO 最低元素: 改为 Pb(0.224), Al(0.246), Os(0.275)')
print('  §3.4 SR SNR: 2.61→3.40 (formation), 0.94→1.39 (hull)')
print('  §3.4 SR R²: ~0.43→~0.44 (10 seeds 平均)')
print('  §3.4 候选: 删 NdCoO3/BiAlO3 (不在实际候选), 改为实际 9 个')
print()
print('=== 合并的节 ===')
print('  原 3.3 (AD) + 3.4 (LOEO) → 新 3.3 (AD and extrapolation)')
print('  原 3.5 (SR) + 3.6 (candidates) → 新 3.4 (Interpretability and screening)')
