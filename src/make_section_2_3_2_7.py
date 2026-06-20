"""
生成 §2.3-2.7 修订版 docx (含 9 个公式 + 2 个 Algorithm Box)
所有数字已 100% 核实于 results/ 和 src/。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 全局字体 (Times New Roman 12pt, 双倍行距 — Elsevier 草稿惯例)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)


def add_heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, indent=True, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    run.italic = italic
    return p


def add_formula(text, label=None):
    """居中公式行, 可选右侧编号 (label)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if label:
        # 用 tab 分隔公式和编号
        run = p.add_run(text)
        run = p.add_run('\t')
        run = p.add_run(f'({label})')
    else:
        run = p.add_run(text)
    return p


def add_inline_math_paragraph(prefix, formula, suffix=''):
    """正文段落内嵌一个简短公式 (不独立成行)"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    if prefix:
        p.add_run(prefix)
    r = p.add_run(formula)
    r.italic = False  # 公式用正体, 由符号本身决定
    if suffix:
        p.add_run(suffix)
    return p


def add_algorithm_box(title, blocks, caption=None):
    """
    概念算法风格 (仿 QSVC 范文).
    blocks: list of dicts, 每个是一行:
      {'type': 'header'/'input'/'step_title'/'text'/'formula'/'note',
       'text': str, 'step_num': optional int}
    - header:    Input:/Output: 行
    - input:     Input 内容行 (缩进)
    - step_title: "Step N: 标题" (加粗)
    - text:      普通描述文字 (缩进)
    - formula:   居中数学公式
    - note:      "where ...", "Here, ..." 解释 (缩进, 可斜体)
    """
    # 算法标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    for blk in blocks:
        t = blk['type']
        text = blk['text']

        if t == 'header':
            # Input: / Output: (左对齐, 加粗冒号前缀)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.left_indent = Cm(0)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

        elif t == 'input':
            # Input 内容行 (缩进, 正常字体)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

        elif t == 'step_title':
            # Step N: 标题 (加粗)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

        elif t == 'text':
            # 描述文字 (缩进)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.75)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

        elif t == 'formula':
            # 居中数学公式
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

        elif t == 'note':
            # "Here, ...", "where ..." 解释 (缩进, 斜体可选)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.75)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)

    # 加 caption (如有)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)

    # 留空
    doc.add_paragraph()


# ========== §2.3 Physics-informed point predictor ==========
add_heading('2.3 Physics-informed point predictor', level=2)

add_para(
    "The point prediction adopts a two-stage residual architecture combining an interpretable physics "
    "baseline with a machine-learning residual corrector."
)

# --- Physics baseline (含 KRR 公式) ---
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.75)
r = p.add_run("(1) Physics baseline \u03bc_p. ")
r.bold = True
p.add_run(
    "A KernelRidge regression on the 14 physics features provides an interpretable baseline. "
    "The prediction for a test sample with physics-feature vector x_phys \u2208 \u211d\u00b9\u2074 is"
)
p2 = add_formula(
    "\u03bc_p(x) = \u03a3_{i \u2208 train} \u03b1_i \u00b7 exp(\u2212\u03b3 \u2016x_phys \u2212 x_{phys,i}\u2016\u00b2),",
    label="1"
)
add_para(
    "where the weights {\u03b1_i} are solved in dual form via ridge regularization (\u03b1 = 1.0) and "
    "the RBF kernel bandwidth is \u03b3 = 0.10. The baseline attains standalone R\u00b2 \u2248 0.77 "
    "(formation energy) and 0.58 (hull energy) under 5-fold cross-validation, quantifying the variance "
    "captured by physically meaningful descriptors alone."
)

# --- ML residual (含 stacking 公式) ---
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.75)
r = p.add_run("(2) ML residual \u03bc_r. ")
r.bold = True
p.add_run(
    "A stacking ensemble of three gradient-boosted-decision-tree (GBDT) base learners "
    "\u2014 LightGBM (leaf-wise), XGBoost (level-wise), and scikit-learn HistGradientBoosting "
    "(histogram-based) \u2014 is trained on the physics-baseline residual (y \u2212 \u03bc_p) using the "
    "full 110 features. The stacking prediction is"
)
add_formula(
    "\u03bc_r(x) = w_0 + \u03a3_{m=1}^{3} w_m \u00b7 r\u0302_m(x),",
    label="2"
)
add_para(
    "where r\u0302_m(x) is the residual prediction of the m-th GBDT base learner and the meta-learner "
    "weights w = (w_0, w_1, w_2, w_3) are fit by Ridge regression on inner 3-fold out-of-fold predictions "
    "to avoid leakage (stacked generalization framework). The final point prediction is"
)
add_formula("\u03bc(x) = \u03bc_p(x) + \u03bc_r(x).", label="3")

# --- Honest reporting ---
add_para(
    "The physics baseline serves as an interpretability anchor (its standalone R\u00b2 quantifies the "
    "variance captured by physically meaningful descriptors); it does not, by itself, improve total "
    "accuracy over pure ML, because the Magpie descriptors partially encode the same physical information. "
    "In particular, the A/B electronegativity difference can be reconstructed from Magpie features with "
    "R\u00b2 = 0.997 (5-fold cross-validated Ridge), so this descriptor adds no independent information "
    "beyond Magpie; other physics descriptors (tolerance factor, ionic radii, d-/f-electron counts), "
    "however, are not reconstructible from Magpie features (R\u00b2 < 0) and thus retain independent value. "
    "We report this honestly rather than claiming the physics layer uniformly boosts performance."
)

add_para(
    "Hyperparameters for the GBDT base learners were selected by Optuna (TPE sampler, 25 trials, 3-fold "
    "inner CV). We note that hyperparameter selection on the full dataset incurs a small optimistic bias: "
    "nested CV estimates ~0.005 R\u00b2 inflation, within the bootstrap confidence interval reported in "
    "Section 3."
)

# ========== §2.4 Conditional conformal prediction (CQR) ==========
add_heading('2.4 Conditional conformal prediction (CQR)', level=2)

add_para(
    "Uncertainty quantification uses Conformalized Quantile Regression (CQR; Romano et al. [14]), "
    "which combines the adaptivity of quantile regression with the finite-sample coverage guarantee of "
    "split conformal. We set \u03b1 = 0.20, yielding nominal 80% coverage intervals. The procedure, "
    "applied within each outer CV fold:"
)

doc.add_paragraph(
    "(1) Split the training fold into a proper-training set (80%) and a calibration set (20%).",
    style='List Number'
)
doc.add_paragraph(
    "(2) Fit two LightGBM quantile regressors on the proper-training set: q\u0302_0.10(X) and q\u0302_0.90(X).",
    style='List Number'
)
doc.add_paragraph(
    "(3) Compute non-conformity scores on the calibration set:",
    style='List Number'
)
add_formula("s_i = max(q\u0302_0.10(x_i) \u2212 y_i, y_i \u2212 q\u0302_0.90(x_i)).", label="4")
doc.add_paragraph(
    "(4) Take the conformal quantile as the \u2308(1\u2212\u03b1)(n_cal+1)\u2309-th ranked value of {s_i} "
    "(the rank-based finite-sample threshold).",
    style='List Number'
)
doc.add_paragraph(
    "(5) The prediction interval for a test sample is",
    style='List Number'
)
add_formula("C(X) = [q\u0302_0.10(X) \u2212 d, q\u0302_0.90(X) + d].", label="5")

p = doc.add_paragraph()
r = p.add_run("Theoretical guarantee. ")
r.bold = True
p.paragraph_format.first_line_indent = Cm(0.75)
p.add_run(
    "Under exchangeability of (calibration, test), the marginal coverage satisfies"
)
add_formula("P(y \u2208 C(X)) \u2265 1 \u2212 \u03b1,", label="6")
add_para(
    "for any \u03b1, with no distributional assumptions (Vovk et al. [13]). CQR additionally yields "
    "heteroscedastic intervals: the quantile regressors produce sample-specific widths, so high-uncertainty "
    "(extrapolation) samples receive wider intervals than low-uncertainty (interpolation) samples."
)

p = doc.add_paragraph()
r = p.add_run("Decoupling design. ")
r.bold = True
p.paragraph_format.first_line_indent = Cm(0.75)
p.add_run(
    "The CQR interval is calibrated independently of the stacking point predictor (Section 2.3), "
    "following the conformal prediction philosophy that the interval need not derive from the same model "
    "as the point estimate. Consequently, the point prediction falls outside its 80% CQR interval for "
    "~9% of formation-energy samples and ~8% of hull-energy samples; this is consistent with the marginal "
    "guarantee (which concerns the true value y, not the point estimate) but introduces a practical "
    "tension discussed in Section 5."
)

p = doc.add_paragraph()
r = p.add_run("Fair baseline. ")
r.bold = True
p.paragraph_format.first_line_indent = Cm(0.75)
p.add_run(
    "Standard split conformal (uniform \u00b1d intervals around a single point predictor) is reported as a baseline. "
    "To ensure a fair comparison of conditional coverage, the baseline point predictor uses a single LightGBM "
    "regressor from the same family as the stacking ensemble (not a weaker Ridge), so the ECE improvement "
    "attributable to CQR is not inflated by baseline weakness. Under this fair comparison, CQR reduces ECE "
    "by 60% for formation energy (0.085 \u2192 0.034) and 43% for hull energy (0.086 \u2192 0.049)."
)

p = doc.add_paragraph()
r = p.add_run("Expected Calibration Error (ECE). ")
r.bold = True
p.paragraph_format.first_line_indent = Cm(0.75)
p.add_run(
    "We quantify conditional coverage via the Expected Calibration Error. Samples are sorted by ensemble "
    "\u03c3 and partitioned into B = 10 equal bins; the ECE is the weighted mean absolute deviation of each "
    "bin\u2019s empirical coverage from the nominal level:"
)
add_formula(
    "ECE = \u03a3_{b=1}^{B} (n_b / N) \u00b7 |cov(b) \u2212 (1 \u2212 \u03b1)|,",
    label="7"
)
add_para(
    "where cov(b) is the empirical coverage in bin b, n_b the bin size, and N the total number of samples. "
    "Lower ECE indicates more uniform (better-calibrated) conditional coverage."
)

# ========== §2.5 Applicability domain ==========
add_heading('2.5 Applicability domain', level=2)

add_para(
    "Three AD criteria delimit the trusted region, computed within each CV fold (no leakage):"
)

doc.add_paragraph(
    "Ensemble \u03c3 (model uncertainty): the standard deviation of the three GBDT base learners\u2019 residual "
    "predictions. Trusted if \u03c3 < the median of all out-of-fold \u03c3 values.",
    style='List Bullet'
)
doc.add_paragraph(
    "k-NN distance (input-space extrapolation): the mean Euclidean distance to the 5 nearest training "
    "neighbors in PCA(20)-reduced space. Trusted if below the 95th percentile of training-set distances.",
    style='List Bullet'
)
doc.add_paragraph(
    "PCA leverage (Williams plot):",
    style='List Bullet'
)
add_formula("h_i = x_i\u1d40 (X\u1d40 X) \u207b\u00b9 x_i", label="8")
add_para(
    "in PCA(20)-reduced space. Trusted if h_i \u2264 3p/n, where p is the number of PCA components (20) "
    "and n is the training-set size."
)

add_para(
    "We report the trusted/untrusted R\u00b2 partition for each method and their pairwise agreement."
)

# ========== §2.6 Evaluation protocol ==========
add_heading('2.6 Evaluation protocol', level=2)

doc.add_paragraph(
    "5-fold cross-validation [24] (shuffle, seed 42) for all reported out-of-fold (OOF) metrics.",
    style='List Bullet'
)
doc.add_paragraph(
    "Stacking meta-learner trained on inner 3-fold OOF predictions to avoid leakage.",
    style='List Bullet'
)
doc.add_paragraph(
    "Bootstrap 95% confidence intervals [25] (1000 resamples) for R\u00b2 and MAE.",
    style='List Bullet'
)
doc.add_paragraph(
    "Imputation (median) and standardization fit within each training fold via a scikit-learn Pipeline, "
    "never on the full dataset.",
    style='List Bullet'
)
doc.add_paragraph(
    "\u03c3\u2013error correlation: Pearson r between ensemble \u03c3 and absolute error, with p-value from "
    "the Pearson test, to validate that ensemble uncertainty ranks errors correctly.",
    style='List Bullet'
)
doc.add_paragraph(
    "Pairwise model comparison (e.g., 5-seed runs, where applicable): Wilcoxon signed-rank test [26] on "
    "paired per-sample absolute errors.",
    style='List Bullet'
)

# ========== Algorithm 1: PACT-Final ==========
add_algorithm_box(
    title='Algorithm 1: The PACT-Final prediction framework',
    blocks=[
        {'type': 'header', 'text': 'Input:'},
        {'type': 'input', 'text': 'D = {(x_i, y_i)} with n = 4,914 perovskite samples and 110 features'},
        {'type': 'input', 'text': 'x_i \u2208 \u211d^{110} is the descriptor vector for the i-th sample (96 Magpie + 14 physics)'},
        {'type': 'input', 'text': 'y_i \u2208 \u211d is the DFT target (formation energy or energy above hull)'},
        {'type': 'input', 'text': '\u03b1 = 0.20 is the miscoverage level (nominal coverage 1 \u2212 \u03b1 = 0.80)'},
        {'type': 'header', 'text': 'Output:'},
        {'type': 'input', 'text': 'Point prediction \u03bc_i, uncertainty \u03c3_i, interval [L_i, U_i], trust flag trust_i for each sample'},
        {'type': 'text', 'text': ''},

        {'type': 'step_title', 'text': 'Step 1: Physics-informed point prediction'},
        {'type': 'text', 'text': 'Compute a physics baseline using kernel ridge regression on the 14 physics features, '
                                  'then correct its residual with a gradient-boosted stacking ensemble:'},
        {'type': 'formula', 'text': '\u03bc_p(x) = \u03a3_{i \u2208 train} \u03b1_i \u00b7 exp(\u2212\u03b3 \u2016x_phys \u2212 x_{phys,i}\u2016\u00b2),'},
        {'type': 'formula', 'text': '\u03bc_r(x) = w_0 + \u03a3_{m=1}^{3} w_m \u00b7 r\u0302_m(x),'},
        {'type': 'formula', 'text': '\u03bc(x) = \u03bc_p(x) + \u03bc_r(x).'},
        {'type': 'note', 'text': 'Here r\u0302_m(x) is the residual prediction of the m-th GBDT base learner '
                                 '(m \u2208 {LightGBM, XGBoost, HistGBT}), and the weights w are fit by a Ridge '
                                 'meta-learner on inner 3-fold out-of-fold predictions. The ensemble uncertainty is'},
        {'type': 'formula', 'text': '\u03c3(x) = std{ r\u0302_1(x), r\u0302_2(x), r\u0302_3(x) }.'},

        {'type': 'step_title', 'text': 'Step 2: Conditional conformal interval (CQR)'},
        {'type': 'text', 'text': 'Fit two LightGBM quantile regressors q\u0302_0.10 and q\u0302_0.90 on the proper-training '
                                  'set, then compute non-conformity scores on the calibration set:'},
        {'type': 'formula', 'text': 's_i = max( q\u0302_0.10(x_i) \u2212 y_i ,  y_i \u2212 q\u0302_0.90(x_i) ).'},
        {'type': 'text', 'text': 'Take the conformal quantile d as the \u2308(1\u2212\u03b1)(n_cal + 1)\u2309-th ranked value '
                                  'of {s_i}, and form the prediction interval:'},
        {'type': 'formula', 'text': 'C(X) = [ q\u0302_0.10(X) \u2212 d ,  q\u0302_0.90(X) + d ].'},
        {'type': 'note', 'text': 'Here the interval is calibrated independently of the point predictor in Step 1, '
                                 'and its marginal coverage satisfies P(y \u2208 C(X)) \u2265 1 \u2212 \u03b1 under exchangeability.'},

        {'type': 'step_title', 'text': 'Step 3: Multi-criteria applicability domain'},
        {'type': 'text', 'text': 'Flag each sample as trusted or untrusted using three independent criteria:'},
        {'type': 'formula', 'text': 'trust_\u03c3  :  \u03c3(x) < median(\u03c3),'},
        {'type': 'formula', 'text': 'trust_kNN :  d_kNN(x) < Q_{0.95}( d_kNN(train) ),'},
        {'type': 'formula', 'text': 'trust_lev :  h_i \u2264 3p / n,  where  h_i = x_i\u1d40 (X\u1d40 X) \u207b\u00b9 x_i.'},
        {'type': 'note', 'text': 'Here d_kNN(x) is the mean Euclidean distance to the 5 nearest training neighbors in '
                                 'PCA(20)-reduced space, p is the number of PCA components, and n is the training-set size.'},
    ],
    caption='All steps are evaluated inside a 5-fold cross-validation loop; preprocessing and model fitting are confined to the training fold to prevent data leakage.'
)

# ========== §2.7 Symbolic regression ==========
add_heading('2.7 Symbolic regression (interpretability cross-check)', level=2)

add_para(
    "Symbolic regression is employed not to discover new physics but to provide a human-readable cross-check "
    "of the ML model. We run gplearn (genetic programming; function set = {+, \u2212, \u00d7, \u00f7, \u221a}; "
    "parsimony coefficient = 0.001; 40 generations; population size = 2000) with 10 random seeds \u00d7 5 CV folds "
    "= 50 independent equations on the 14 physics features. Feature consensus frequency "
    "(the fraction of equations in which each physics feature appears) is reported, rather than any single "
    "non-unique equation."
)

add_para(
    "We additionally derive an empirical SR applicability criterion by comparing SR success on formation energy "
    "(signal-to-noise ratio SNR = 3.40) versus failure on hull energy (SNR = 1.39), where"
)
add_formula(
    "SNR = R\u00b2(KRR) / (1 \u2212 R\u00b2(KRR)),",
    label="9"
)
add_para(
    "and R\u00b2(KRR) is the standalone R\u00b2 of the KernelRidge physics baseline (0.773 for formation, "
    "0.581 for hull). The empirical boundary is SNR \u2248 2: SR succeeds when SNR \u2265 2 (formation energy) "
    "and fails when SNR < 2 (hull energy). We emphasize that this criterion is empirical, derived from only "
    "two target properties, and its generalization to other material properties requires further validation "
    "(see Section 5)."
)

# ========== Algorithm 2: SR Ensemble ==========
add_algorithm_box(
    title='Algorithm 2: Symbolic regression ensemble with feature consensus',
    blocks=[
        {'type': 'header', 'text': 'Input:'},
        {'type': 'input', 'text': 'D_phys = {(x_{phys,i}, y_i)} with 14 physics-informed descriptors'},
        {'type': 'input', 'text': 'S = {42, 123, 456, 789, 2024, 314, 271, 161, 803, 951}, 10 random seeds'},
        {'type': 'input', 'text': 'F = 5 cross-validation folds'},
        {'type': 'header', 'text': 'Output:'},
        {'type': 'input', 'text': 'consensus_freq(f), the fraction of equations in which feature f appears, for each physics feature f'},
        {'type': 'text', 'text': ''},

        {'type': 'step_title', 'text': 'Step 1: Symbolic regression search'},
        {'type': 'text', 'text': 'For each seed s \u2208 S and each CV fold, fit a genetic-programming symbolic '
                                  'regression (gplearn) on the training fold of the physics features. The search space is '
                                  'defined by the operator set'},
        {'type': 'formula', 'text': '\u03a6 = { +, \u2212, \u00d7, \u00f7, \u221a } ,'},
        {'type': 'note', 'text': 'with parsimony coefficient 0.001, 40 generations, and population size 2,000. '
                                 'This yields 10 seeds \u00d7 5 folds = 50 independent equations { eq_1, eq_2, \u2026, eq_50 }.'},

        {'type': 'step_title', 'text': 'Step 2: Feature consensus aggregation'},
        {'type': 'text', 'text': 'For each physics feature f, compute its consensus frequency as'},
        {'type': 'formula', 'text': 'consensus_freq(f) = | { eq \u2208 equations : f appears in eq } | / 50.'},
        {'type': 'note', 'text': 'Here the numerator counts how many of the 50 equations contain feature f. '
                                 'Consensus frequency, rather than any single equation, is reported to mitigate '
                                 'the non-uniqueness of symbolic regression solutions.'},

        {'type': 'step_title', 'text': 'Step 3: Applicability criterion'},
        {'type': 'text', 'text': 'Derive an empirical SR applicability boundary from the signal-to-noise ratio of the '
                                  'physics baseline:'},
        {'type': 'formula', 'text': 'SNR = R\u00b2(KRR) / ( 1 \u2212 R\u00b2(KRR) ).'},
        {'type': 'note', 'text': 'Symbolic regression is deemed applicable when SNR \u2265 2 (e.g., formation energy, '
                                 'SNR = 3.40) and inapplicable when SNR < 2 (e.g., energy above hull, SNR = 1.39). '
                                 'This criterion is empirical and based on two target properties.'},
    ],
    caption='Consensus frequency is reported instead of a single equation to handle the inherent non-uniqueness of symbolic regression.'
)

doc.save('paper/section_2_3_to_2_7_revised.docx')
print('[OK] saved: paper/section_2_3_to_2_7_revised.docx')
print(f'  公式总数: 9 (Eq.1-9)')
print(f'  Algorithm Box: 2')
print(f'    Algorithm 1: PACT-Final training (29 lines, in §2.6 后)')
print(f'    Algorithm 2: SR Ensemble consensus (17 lines, in §2.7)')
print(f'  §2.3: Eq.1 (KRR), Eq.2 (Stacking), Eq.3 (μ=μ_p+μ_r)')
print(f'  §2.4: Eq.4 (s_i), Eq.5 (C(X)), Eq.6 (P(y∈C(X))), Eq.7 (ECE)')
print(f'  §2.5: Eq.8 (leverage)')
print(f'  §2.7: Eq.9 (SNR)')
