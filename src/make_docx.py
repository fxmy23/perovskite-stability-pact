"""
Generate Elsevier-formatted .docx from manuscript.md + tables.md + references.md
CMS format: single column, Times New Roman 12pt, double-spaced, numbered sections.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

PAPER = Path('paper')
doc = Document()

# === Page setup (A4, 1 inch margins) ===
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# === Default style: Times New Roman 12pt, double-spaced ===
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 2.0

def add_heading_custom(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = 'Times New Roman'
    return p

def add_para(text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# Read manuscript
md = (PAPER / 'manuscript.md').read_text(encoding='utf-8')
lines = md.split('\n')

i = 0
in_table = False
while i < len(lines):
    line = lines[i].strip()
    # Skip markdown artifacts
    if line.startswith('Target journal:'):
        i += 1; continue
    if line == '---':
        i += 1; continue
    # Title (H1)
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        i += 1; continue
    # Author/Affiliation (bold lines after title)
    if line.startswith('**Author:**') or line.startswith('**Affiliation:**') or line.startswith('**Corresponding'):
        text = line.replace('**', '')
        add_para(text, bold=False)
        i += 1; continue
    # Section headings
    if line.startswith('## '):
        heading = line[3:].strip()
        if heading.lower() not in ['abstract', 'keywords']:
            add_heading_custom(heading, 1)
        i += 1; continue
    if line.startswith('### '):
        heading = line[4:].strip()
        add_heading_custom(heading, 2)
        i += 1; continue
    # Keywords/Abstract inline
    if line.startswith('**Keywords:**'):
        add_para(line.replace('**', ''), italic=False)
        i += 1; continue
    # Regular paragraph (skip markdown tables, figures placeholders, empty lines)
    if line and not line.startswith('|') and not line.startswith('[**') and not line.startswith('[**Figure'):
        # Clean markdown
        clean = line.replace('**', '').replace('*', '')
        if clean.strip():
            add_para(clean)
    i += 1

# === Add Tables section ===
doc.add_page_break()
add_heading_custom('Tables', 1)
tables_md = (PAPER / 'tables.md').read_text(encoding='utf-8')
for tline in tables_md.split('\n'):
    tline = tline.strip()
    if tline.startswith('## '):
        add_heading_custom(tline[3:].strip(), 2)
    elif tline.startswith('|') and '---' not in tline:
        # Parse table row
        cells = [c.strip() for c in tline.split('|')[1:-1]]
        if not hasattr(add_para, '_table_started') or not add_para._table_started:
            doc.add_paragraph()  # spacer
            add_para._table_started = True
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(' | '.join(cells))
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    elif tline.startswith('Note:'):
        add_para(tline, italic=True)

# === Add References ===
doc.add_page_break()
add_heading_custom('References', 1)
refs_md = (PAPER / 'references.md').read_text(encoding='utf-8')
for rline in refs_md.split('\n'):
    rline = rline.strip()
    if rline.startswith('['):
        add_para(rline)
    elif rline.startswith('#'):
        pass  # skip headers
    elif rline.startswith('---') or rline.startswith('## ') or 'Notes' in rline:
        pass

# === Add Figure captions ===
doc.add_page_break()
add_heading_custom('Figure Captions', 1)
captions = [
    'Fig. 1. Overview of the uncertainty-aware prediction framework (PACT-Final). Input features feed a physics-informed point predictor (KernelRidge baseline + GBDT stacking residual) and, independently, a conditional conformal interval (CQR) plus applicability domain (sigma/k-NN/leverage). Output includes point estimate, adaptive 80% interval, and trust label.',
    'Fig. 2. Parity plots of predicted vs. DFT values for (a) formation energy and (b) energy above hull. Points colored by ensemble sigma. Dashed line: y = x. R-squared and MAE annotated.',
    'Fig. 3. Reliability diagrams showing empirical coverage per uncertainty decile for standard split conformal vs. CQR. Dotted line: nominal 0.80 coverage. CQR tracks the nominal level more uniformly.',
    'Fig. 4. Applicability domain visualization. Predicted value vs. absolute error, colored by trusted (sigma < median) vs. untrusted regions.',
    'Fig. 5. Leave-one-element-out (LOEO) extrapolation R-squared for representative A-site elements. Pure ML vs. SR + ML.',
    'Fig. 6. (a) LightGBM SHAP feature importance (top-15) and (b) symbolic regression consensus frequency for physics descriptors.',
    'Graphical Abstract. Vertical workflow summarizing the framework from input perovskite compounds to reliability-annotated stability predictions.',
]
for cap in captions:
    add_para(cap)

# Save
output = PAPER / 'manuscript_cms_v2.docx'
doc.save(str(output))
print(f'[SAVE] {output} ({output.stat().st_size // 1024} KB)')
