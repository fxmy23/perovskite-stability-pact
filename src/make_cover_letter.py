"""
Generate cover_letter.docx for CMS submission.
- Clean of all encoding errors (no 鈧?/R虏/蟽 garbage)
- Numbers updated to verified values (73 elements, R²=0.739, SNR 3.40/1.39)
- Follows Elsevier official structure
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)


def add_para(text, bold=False, indent=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_heading_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p


# === Header ===
add_heading_line('Cover Letter \u2014 Computational Materials Science submission')
add_para('To: The Editor-in-Chief, Computational Materials Science')
add_para('Date: [submission date]')
add_para('Subject: Submission of manuscript \u201cUncertainty-Aware Prediction of '
         'Perovskite Oxide Stability with Conditional Conformal Intervals and '
         'Applicability-Domain Guidance\u201d')
add_para('Article type: Full research article (single author)')
add_para('')

# === Salutation ===
add_para('Dear Editor,')

# === Para 1: who I am ===
add_para(
    'I am pleased to submit my manuscript entitled \u201cUncertainty-Aware Prediction '
    'of Perovskite Oxide Stability with Conditional Conformal Intervals and '
    'Applicability-Domain Guidance\u201d for consideration as a full research article '
    'in Computational Materials Science. I am a junior undergraduate student at '
    'Weiyang College, Tsinghua University, and this work was conducted as an '
    'independent research project using open data and a personal workstation.'
)

# === Para 2: the gap ===
add_para('The gap this work addresses. ', bold=False)
p = doc.paragraphs[-1]
p.add_run(
    'Machine learning has become an indispensable tool for accelerating the '
    'discovery of ABO\u2083 perovskite oxides, with recent descriptor-based models '
    'reaching R\u00b2 > 0.91 for formation-energy prediction. However, the vast '
    'majority of published models report only point predictions and aggregate '
    'accuracy metrics, without quantifying when their predictions can be trusted '
    '\u2014 a critical gap for high-throughput screening pipelines where '
    'false-positive stability predictions waste expensive DFT verification budget. '
    'This manuscript addresses that gap directly by integrating point prediction, '
    'calibrated uncertainty, applicability-domain guidance, and extrapolation '
    'assessment into a single reproducible framework.'
)
p.runs[0].bold = True

# === Para 3: contributions heading ===
add_para('What this work contributes to computational materials science:', bold=True)

contributions = [
    ('A reliability-aware prediction framework, not just an accuracy benchmark. ',
     'I integrate a physics-informed point predictor (KernelRidge baseline + GBDT '
     'stacking residual) with Conformalized Quantile Regression (CQR), which '
     'provides sample-adaptive prediction intervals satisfying a finite-sample, '
     'distribution-free coverage guarantee. Under a fair same-family baseline, '
     'CQR reduces the Expected Calibration Error by 43\u201360%, yielding intervals '
     'that widen appropriately for extrapolation samples and tighten for '
     'interpolation samples.'),
    ('An applicability domain for perovskite stability models. ',
     'Three complementary applicability-domain criteria \u2014 ensemble \u03c3, '
     'k-NN distance, and PCA leverage \u2014 are compared, and the trusted region '
     '(R\u00b2 = 0.945 for formation energy) is clearly separated from the '
     'untrusted region (R\u00b2 = 0.886), giving experimentalists a concrete '
     'boundary for reliable use.'),
    ('Element-level extrapolation mapping. ',
     'A leave-one-element-out evaluation over all 73 A/B-site elements quantifies '
     'where the model generalizes and where it fails (mean per-element '
     'R\u00b2 = 0.739, no negative-R\u00b2 elements), directly informing screening '
     'of unseen chemistries.'),
    ('Interpretability grounded in materials chemistry. ',
     'A symbolic-regression ensemble (50 independent equations) recovers the known '
     'dominance of A-site electronegativity in formation energy, and an empirical '
     'applicability boundary for symbolic regression is derived by contrasting '
     'formation energy (SNR = 3.40, SR succeeds) with hull energy (SNR = 1.39, '
     'SR fails).'),
]
for i, (lead, body) in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f'{i}. ')
    r1.bold = True
    r2 = p.add_run(lead)
    r2.bold = True
    p.add_run(body)

# === Para 4: fit with scope ===
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Fit with the journal\u2019s scope. ')
r.bold = True
p.add_run(
    'I believe this work fits the scope of Computational Materials Science '
    '\u2014 advancing computational methods applied to materials prediction '
    '\u2014 and would be of interest to readers working on ML-accelerated '
    'materials discovery, uncertainty quantification, and perovskite design. I '
    'emphasize that the value of this work lies in the uncertainty, '
    'applicability-domain, and extrapolation components, which make high-'
    'throughput screening more trustworthy, rather than in chasing the highest '
    'R\u00b2.'
)

# === Para 5: reproducibility ===
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Reproducibility and transparency. ')
r.bold = True
p.add_run(
    'All source code and processed data are provided under the MIT license at '
    'https://github.com/fxmy23/perovskite-stability-pact, with the main results '
    'reproducible via a single command (python src/pact_final.py). The limitations '
    'of the framework \u2014 single data source, decoupling tension between point '
    'and interval predictions, absence of DFT-validated candidates \u2014 are '
    'stated explicitly in the manuscript, which I believe strengthens rather '
    'than weakens the contribution.'
)

# === Para 6: declarations ===
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Declarations. ')
r.bold = True
p.add_run(
    'This manuscript has not been published previously, is not under consideration '
    'for publication elsewhere, and all its data and results are original to this '
    'work. I declare no conflicts of interest. The use of AI-assisted tools during '
    'manuscript preparation is declared in the manuscript, and I take full '
    'responsibility for its content.'
)

# === Closing ===
add_para('Thank you for your consideration. I look forward to your response.')
add_para('')
add_para('Sincerely,')
add_para('')
p = doc.add_paragraph()
r = p.add_run('Xumingyong Feng')
r.bold = True
add_para('Weiyang College, Tsinghua University, Beijing 100084, China')
add_para('Email: fxmy23@mails.tsinghua.edu.cn')

doc.save('paper/cover_letter.docx')
print('[OK] saved: paper/cover_letter.docx')

# Quick sanity checks
import re
full = '\n'.join(p.text for p in doc.paragraphs)
# Check no encoding garbage
garbage = re.findall(r'[鈧虏蟽\xa0]', full)
print(f'Encoding garbage chars: {len(garbage)} -> {"OK" if len(garbage)==0 else "STILL PRESENT"}')
# Check key numbers
checks = [
    ('73 elements', '73' in full),
    ('R\u00b2 = 0.739', '0.739' in full),
    ('SNR = 3.40', '3.40' in full),
    ('SNR = 1.39', '1.39' in full),
    ('43\u201360%', '43' in full and '60%' in full),
    ('0.945 trusted', '0.945' in full),
]
for name, ok in checks:
    print(f'  {"\u2713" if ok else "\u2717"} {name}')
# Check for old/wrong numbers
print('Old-number check:')
print(f'  "68 elements" present: {"68" in full}  (should be False)')
print(f'  "R\u00b2 averages 0.70" present: {"0.70" in full}  (should be False)')
print(f'  "SNR = 2.61" present: {"2.61" in full}  (should be False)')
print(f'  "SNR = 0.94" present: {"0.94" in full}  (should be False)')
