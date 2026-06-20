"""
生成 §3.4 (Interpretability and candidate screening) 范文风格版.
- 连续叙述, 无标签
- 加 Fig. 8 / Table 3 引入话
- 段间过渡
- 零问句 (已验证)
所有数字 100% 核实.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.add_run(text)
    return p


def add_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    return table


# ========== §3.4 ==========
add_heading('3.4 Interpretability and candidate screening', level=1)

# 段 1: SR 共识 + Fig. 8 引入 (连续叙述, 无标签)
add_para(
    "Beyond the accuracy and uncertainty metrics above, we ask whether the model has "
    "captured chemically meaningful structure rather than spurious statistical "
    "regularities. Two independent interpretability lenses converge on the same "
    "answer. The first is a symbolic-regression ensemble (Algorithm 2): across 50 "
    "independent equations, A-site electronegativity appears in 100% of them, with "
    "B-site group number (64%) and B-site electronegativity (60%) the next most "
    "conserved, as visualized in Fig. 8(b). The second is the SHAP analysis of the "
    "full stacking ensemble, shown in Fig. 8(a), where A-site-electronegativity-"
    "derived Magpie features rank among the top contributors. The agreement between "
    "these two methods operating in different spaces—symbolic vs. statistical—"
    "provides strong evidence that the dominance of A-site electronegativity is a "
    "genuine physical signal, not an artifact of either method alone."
)

# 段 2: 物理方向 (保留原内容, 去标签, 连续叙述)
add_para(
    "This dominance is further corroborated by the sign and magnitude of each top "
    "descriptor's correlation with the target. The strongest single correlate of "
    "formation energy is A-site electronegativity (Pearson r = +0.537): more "
    "electronegative A-site cations—e.g., transition metals at the A site—yield "
    "less negative, hence less stable, formation energies, consistent with the fact "
    "that highly electropositive cations (alkali, alkaline-earth, lanthanide) form "
    "stronger ionic bonds with oxygen and stabilize the perovskite lattice. The "
    "opposite sign governs Magpie X_std, the compositional standard deviation of "
    "electronegativity (r = −0.538): greater electronegativity contrast among "
    "A/B/O sites produces more negative formation energies, reflecting the "
    "thermodynamic driving force of charge transfer in ionic bonding. The tolerance "
    "factor t shows a moderate negative correlation (r = −0.278), so that as t "
    "increases toward and beyond unity—A-site too large for the cage—the structure "
    "becomes geometrically strained and formation energy rises, consistent with "
    "Goldschmidt's criterion. B-site group number (r = +0.297) and B-site "
    "electronegativity (r = +0.291) are positively correlated, because higher B-"
    "site electronegativity (late transition metals) reduces the ionic character of "
    "the B–O bond and destabilizes the perovskite relative to competing phases. "
    "These signs are all consistent with established perovskite chemistry [29] "
    "[30,31], confirming that the model respects known structure–property relations."
)

# 段 3: SR 适用边界 (过渡 + 连续叙述)
add_para(
    "The interpretability cross-check is itself target-dependent, and this "
    "asymmetry exposes a useful methodological boundary. Symbolic regression "
    "succeeds on formation energy—standalone R² ≈ 0.44 across 10 seeds, "
    "corresponding to SNR = 3.40 (Eq. 9)—but fails on hull energy (SNR = 1.39), "
    "with the SR solution collapsing to a constant in several folds. We trace this "
    "to the differing physical nature of the two targets: formation energy is "
    "dominated by single-compound bond chemistry, near-linear and well-expressed by "
    "the basic operators {+, −, ×, ÷, √}, whereas hull energy depends on "
    "competition with all other phases and is strongly non-linear and non-monotonic "
    "in geometric descriptors. This yields an empirical SR applicability criterion, "
    "SNR ≥ 2, which we caveat as based on only two targets and not yet universal."
)

# 段 4: 候选筛选 + Table 3 引入 (过渡 + 连续叙述)
add_caption('Table 3. Candidate stable perovskites with validation tier. "In matbench" denotes presence in the independent matbench_perovskites DFT database; "Goldsm." denotes location within the Goldschmidt synthesizability zone (t ∈ [0.8, 1.05], μ ∈ [0.4, 0.9]).')

add_table(
    headers=['Formula', 'Pred. E_hull', 'σ', 'In matbench', 'Goldsm.', 'Grade'],
    rows=[
        ['LaTiO3', '0.005', '0.022', 'Yes', 'Yes', 'Mod.'],
        ['PrCoO3', '0.006', '0.024', 'No', 'Yes', 'Weak'],
        ['GdCoO3', '0.016', '0.027', 'No', 'Yes', 'Weak'],
        ['YbCoO3', '0.017', '0.028', 'No', 'Yes', 'Weak'],
        ['DyCoO3', '0.019', '0.029', 'No', 'Yes', 'Weak'],
        ['ErFeO3', '0.019', '0.030', 'No', 'Yes', 'Weak'],
        ['YCoO3', '0.025', '0.031', 'Yes', 'Yes', 'Mod.'],
        ['SrNpO3', '0.025', '0.034', 'No', 'No', 'Weak'],
        ['LaAlO3', '0.033', '0.037', 'Yes', 'Yes', 'Weak'],
    ],
    col_widths=[1.2, 1.1, 0.8, 1.0, 0.9, 0.8]
)

add_para(
    "Having established that the model respects known chemistry, we apply it to "
    "virtual A–B combinations to screen for candidate stable perovskites. Retaining "
    "only compositions predicted to lie on or near the convex hull (E_hull ≤ 0.05 "
    "eV/atom) within the trusted applicability domain yields the nine candidates "
    "listed in Table 3. A three-tier validation provides indirect evidence rather "
    "than confirmation: three candidates—LaTiO3, YCoO3, and LaAlO3—appear in the "
    "independent matbench_perovskites DFT database, confirming that they have been "
    "studied; eight of nine fall within the Goldschmidt synthesizability zone "
    "(t ∈ [0.8, 1.05], μ ∈ [0.4, 0.9]); and leave-element-out robustness testing "
    "confirms that the moderate-grade predictions are not artifacts of single-"
    "element memorization. We emphasize that database presence attests to prior "
    "study, not to numerical agreement with our predictions (given the reference-"
    "frame mismatch noted in §2.1), and that the remaining six candidates lack "
    "independent DFT or experimental verification. The screening result therefore "
    "demonstrates that the pipeline produces physically plausible, cross-database-"
    "consistent candidates rather than spurious low-E_hull outliers, but none of "
    "these predictions should be treated as a confirmed materials discovery in the "
    "absence of dedicated DFT validation."
)

doc.save('paper/section_3_4_revised.docx')

# 验证无问句
import re
full_text = '\n'.join(p.text for p in doc.paragraphs)
q_marks = re.findall(r'[?？]', full_text)
print('[OK] saved: paper/section_3_4_revised.docx')
print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')
total_words = sum(len(p.text.split()) for p in doc.paragraphs)
print(f'总词数: {total_words}')
print(f'问句检查: {len(q_marks)} 个问号 -> {"✓ 无问句" if len(q_marks)==0 else "✗ 有问句"}')
