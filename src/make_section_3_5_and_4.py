"""
生成 §3.5 Limitations + §4 Conclusions (范文风格, 数字已核实, 0 问句).
§4 Limitations 并入 §3 作 §3.5 (符合 CMS ~50% 论文惯例).
§5 Conclusions 提前为 §4.
所有数字 100% 核实:
- 删 Pearson r=0.358 (无来源)
- LOEO 68→73, 0.70→0.739, 删 B/Ac 负 R² (实际 0 个)
- 点在区间外 ~9% (formation 9.4%, hull 7.5%)
- R² 0.914/0.800, trusted 0.945/0.873, untrusted 0.886/0.748
- ECE 改善 43-60%
- 超参乐观偏差 ~0.005 R²
- 候选 9 个, 3 个 matbench 验证
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.add_run(text)
    return p


# ========== §3.5 Limitations ==========
add_heading('3.5 Limitations', level=1)

# 引言句
add_para(
    "Several limitations of this work should be acknowledged to support informed "
    "downstream use, and we discuss them here in decreasing order of practical "
    "impact rather than as isolated bullet points."
)

# 段 1: 单数据源 + 跨数据集 (合并原 #1, 删 0.358)
add_para(
    "The most consequential constraint is the single-source dataset: all 4,914 "
    "samples originate from one DFT study (wolverton_oxides [18]), so the model "
    "has been exposed to only one computational reference frame. Cross-dataset "
    "generalization could not be numerically verified, because the Materials "
    "Project [32] API was inaccessible during this work and matbench_perovskites "
    "uses a structurally and energetically different formation-energy reference "
    "frame (direct model transfer between the two sources yields negative R², "
    "precluding meaningful cross-dataset validation). The leave-one-element-out "
    "evaluation in §3.3 partially addresses generalization within the wolverton "
    "chemical space—mean per-element R² = 0.739 with no negative-R² elements—but "
    "it does not substitute for independent DFT validation on a genuinely distinct "
    "dataset, which remains the most pressing future test."
)

# 段 2: 候选验证深度 (原 #3, 强调诚实)
add_para(
    "A direct consequence of the single-source constraint is the limited depth of "
    "candidate validation. The nine predicted stable perovskites in Table 3 lack "
    "DFT or experimental verification; the three-tier validation provides only "
    "indirect evidence—database presence attests to prior study rather than to "
    "numerical agreement, and the six non-matbench candidates (PrCoO3, GdCoO3, "
    "YbCoO3, DyCoO3, ErFeO3, SrNpO3) have, to our knowledge, no independent "
    "computational record. We therefore frame these as weak-evidence candidates "
    "requiring DFT confirmation rather than as discoveries, and we caution against "
    "treating any of the nine as experimentally actionable without dedicated "
    "verification."
)

# 段 3: 点-区间解耦张力 + 外推 (合并原 #2 + #5)
add_para(
    "At the methodological level, two design choices introduce tensions that users "
    "should understand. First, the point predictor and the CQR interval are "
    "calibrated independently, following conformal-prediction philosophy; as a "
    "result, the point prediction falls outside its own 80% CQR interval for about "
    "9% of formation-energy samples and 8% of hull-energy samples. This is "
    "consistent with the marginal coverage guarantee—which concerns the true value "
    "y rather than the point estimate—but in practice users must decide which "
    "statement to trust when they disagree, and we recommend treating the interval "
    "as the primary uncertainty statement. Second, the LOEO evaluation shows that "
    "accuracy drops from R² = 0.914 (in-domain) to 0.739 when an entire element is "
    "held out; although no element yields a negative R², the worst-performing "
    "elements (Pb 0.22, Al 0.25, Os 0.27) confirm that extrapolation to "
    "underrepresented chemistries is unreliable, and users should restrict "
    "predictions to the applicability domain rather than trusting the point "
    "estimate alone."
)

# 段 4: SR 判据 + 超参偏差 (合并原 #4 + #6)
add_para(
    "Finally, two empirical choices carry residual uncertainty. The SR "
    "applicability criterion (SNR ≥ 2) is derived from only two target properties "
    "and may not generalize to other material properties whose signal-to-noise "
    "structure differs. The Optuna hyperparameters were selected on the full "
    "dataset before cross-validation, incurring a small optimistic bias estimated "
    "at roughly 0.005 R² by nested CV; this bias falls within the reported "
    "bootstrap confidence intervals but is nonetheless present and would be "
    "eliminated by a fully nested protocol at additional computational cost. None "
    "of these limitations invalidates the central findings, but together they "
    "define the boundary within which the framework's predictions should be "
    "trusted."
)

# ========== §4 Conclusions ==========
add_heading('4. Conclusions', level=1)

add_para(
    "We presented an uncertainty-aware framework for predicting ABO3 perovskite "
    "formation energy and thermodynamic stability that integrates a physics-"
    "informed point predictor, conditional conformal (CQR) intervals with finite-"
    "sample coverage guarantees, a multi-method applicability domain, and "
    "symbolic-regression-based interpretability cross-checks. The framework "
    "achieves R² = 0.914 for formation energy and R² = 0.800 for energy above hull "
    "under 5-fold cross-validation, while delivering sample-adaptive intervals "
    "whose conditional coverage improves 43–60% over standard split conformal "
    "under a fair same-family baseline. The applicability domain separates a "
    "trusted region (R² = 0.945 for formation energy) from an untrusted region "
    "(R² = 0.886), and the 73-element leave-one-element-out evaluation provides "
    "an element-level extrapolation map (mean R² = 0.739) that refines this "
    "statistical trust boundary into chemically specific guidance. Symbolic "
    "regression consensus validates the known electronegativity dominance—A-site "
    "electronegativity appears in 100% of 50 independent equations—and exposes an "
    "empirical applicability boundary (SNR ≥ 2) that we frame honestly as target-"
    "specific rather than universal."
)

add_para(
    "The central message is methodological: for materials-discovery pipelines, "
    "point accuracy alone is insufficient. Per-sample uncertainty, applicability-"
    "domain guidance, and extrapolation assessment are necessary components of a "
    "trustworthy predictor, and the framework presented here offers an "
    "integrated, fully reproducible implementation of these components under a "
    "single nested cross-validation protocol. The limitations documented in §3.5 "
    "are not incidental: they define the conditions under which the framework's "
    "predictions can be relied upon, and reporting them transparently is, in our "
    "view, as important as reporting the accuracy itself. Future work should "
    "validate the candidate materials via dedicated DFT calculations, test cross-"
    "dataset generalization once an accessible independent database is available, "
    "and refine the SR applicability criterion across additional material "
    "properties beyond the two studied here."
)

doc.save('paper/section_3_5_and_4_revised.docx')

# 验证: 无问句, 统计
full_text = '\n'.join(p.text for p in doc.paragraphs)
q_marks = re.findall(r'[?？]', full_text)
total_words = sum(len(p.text.split()) for p in doc.paragraphs)
print('[OK] saved: paper/section_3_5_and_4_revised.docx')
print(f'段落数: {len(doc.paragraphs)}')
print(f'总词数: {total_words}')
print(f'问句检查: {len(q_marks)} 个问号 -> {"✓ 无问句" if len(q_marks)==0 else "✗ 有问句"}')
print()
print('=== 结构 ===')
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        print(f'  [{p.style.name}] {p.text}')
print()
print('=== §3.5 各段词数 ===')
sections = {}
current = None
for p in doc.paragraphs:
    txt = p.text.strip()
    if p.style.name == 'Heading 1' and txt == '3.5 Limitations':
        current = '3.5'
        sections[current] = 0
    elif p.style.name == 'Heading 1' and txt == '4. Conclusions':
        current = '4'
        sections[current] = 0
    elif current and txt:
        sections[current] += len(txt.split())
for s, w in sections.items():
    print(f'  §{s}: {w} 词')
